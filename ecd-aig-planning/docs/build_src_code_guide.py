from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "SRC_CODE_GUIDE.md"
OUT = ROOT / "outputs" / "ECD-AIG_SRC_CODE_GUIDE.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(92, 102, 112)
INK = RGBColor(33, 37, 41)


def set_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 9),
        ("Heading 2", 13, BLUE, 13, 6),
        ("Heading 3", 11.5, DARK_BLUE, 9, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor(40, 48, 56)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.left_indent = Inches(0.14)
    code.paragraph_format.line_spacing = 1.0

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    bullet.font.size = Pt(10.5)
    bullet.paragraph_format.space_after = Pt(2)
    bullet.paragraph_format.line_spacing = 1.15


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(82)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ECD-AIG Planning")
    set_font(run, size=28, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("src/ecd_aig 파이썬 소스 코드 해설서")
    set_font(run, size=17, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run("응답 전 ECD-AIG 설계 추적과 위험 점검 코드를 읽기 위한 개발자 안내서")
    set_font(run, size=11.5, color=MUTED, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("핵심 원칙")
    set_font(run, size=11, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("사전 점검 통과 != 경험적 타당도 확보")
    set_font(run, size=14, color=RGBColor(155, 28, 28), bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated for the local research prototype | 2026-06-01")
    set_font(run, size=9.5, color=MUTED)
    doc.add_page_break()


def add_index(doc):
    doc.add_heading("파일 빠른 찾기", level=1)
    p = doc.add_paragraph("처음 읽을 때는 models.py → validation.py → item_quality.py → pre_response.py → __main__.py 순서를 권장한다.")
    p.paragraph_format.space_after = Pt(8)
    files = [
        ("핵심 흐름", "models.py, validation.py, item_quality.py, pre_response.py, __main__.py"),
        ("문항 준비", "generation.py, import_items.py, blueprint.py, simulation.py"),
        ("설명 산출물", "ecd_report.py, caf.py, toulmin.py, dossier.py"),
        ("응답자료 이후", "scoring.py, response_data.py, psychometrics.py, agreement.py"),
        ("출력과 화면", "rendering.py, export.py, webapp.py, __init__.py"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1900, 7460])
    table.rows[0].cells[0].text = "범주"
    table.rows[0].cells[1].text = "파일"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            set_font(run, size=9.5, bold=True, color=DARK_BLUE)
    for label, value in files:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for cell in cells:
            for run in cell.paragraphs[0].runs:
                set_font(run, size=9.2)
    doc.add_paragraph()


def add_markdown(doc):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines = []
    first_h1_skipped = False
    for line in lines:
        if line.startswith("```"):
            if in_code:
                for code_line in code_lines:
                    p = doc.add_paragraph(style="Code Block")
                    p.add_run(code_line or " ")
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            continue
        if line.startswith("# "):
            if not first_h1_skipped:
                first_h1_skipped = True
                continue
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line[:3].isdigit() and line[1:3] == ". ":
            doc.add_paragraph(line, style="List Number")
        else:
            p = doc.add_paragraph()
            p.add_run(line)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("ECD-AIG Planning | Source Code Guide")
    set_font(run, size=8.5, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    add_cover(doc)
    add_index(doc)
    add_markdown(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
